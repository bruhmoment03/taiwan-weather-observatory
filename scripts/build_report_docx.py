# -*- coding: utf-8 -*-
"""組裝期末報告 Word 檔：report/cover.json + report/findings.json + assets/report*.png
→ report/期末報告.docx。可重複執行（影片連結補上後重新生成）。"""
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

BASE = Path(__file__).resolve().parents[1]
REPORT = BASE / "report"
ASSETS = BASE / "assets"

FILES_TABLE = [
    ("user_data.csv", "313", "7", "學校、年級、班級、國數英成績"),
    ("dp001_prac.csv", "6,624", "9", "練習作答（正確率、能力指標、作答串）"),
    ("dp001_review.csv", "4,567", "11", "影片瀏覽（影片、起迄、完成率）"),
    ("dp001_review_plus.csv", "118,390", "6", "影片操作事件（播放/暫停/拖曳/筆記）"),
    ("dp001_exam.csv", "1,763", "7", "影片檢核點作答"),
    ("dp002_exam.csv", "41,864", "8", "測驗平臺題項作答紀錄"),
    ("dp003_word.csv", "2,140", "13", "英文單字遊戲（目標單字、累積對錯）"),
    ("dp003_math.csv", "1,365", "10", "數學遊戲（單元、對錯、耗時）"),
    ("dp004_interaction.csv", "15,392", "9", "綜合平臺測驗互動"),
    ("dp004_video.csv", "6,086", "7", "綜合平臺觀影紀錄"),
    ("dp004_webpage.csv", "23,642", "6", "綜合平臺學習資源瀏覽"),
]

RENAME_TABLE = [
    ("user_sn", "使用者編號"), ("organization_id", "學校代碼"), ("grade", "年級"),
    ("class", "班級"), ("chinese_score", "國文成績"), ("math_score", "數學成績"),
    ("english_score", "英語成績"), ("subject_name", "科目"), ("video_name", "影片名稱"),
    ("finish_rate", "完成率"), ("indicator_name", "能力指標"), ("score_rate", "練習正確率"),
]

PREP_STEPS = [
    "Tab 分隔原始檔轉為標準 CSV，統一輸出 UTF-8（含 BOM）以利 Power BI 讀取中文。",
    "欄位重新命名：所有英文欄名改為中文（對照表如表 2），提升報表可讀性。",
    "缺失值處理：國文、數學、英語成績各有 120、45、112 筆缺漏；行為彙總保留全部 313 位使用者，"
    "與成績相關之分析則逐科排除缺漏列並於圖表註記樣本數。",
    "時間標準化：dp002/dp004 之 ISO8601 含時區時間戳與 dp003 之 Unix 毫秒時間戳，統一轉為台北時間，"
    "並衍生年、月、週次、星期、時段欄位以利時間序列分析（dp003_math 無動作時間欄，"
    "以伺服器寫入時間 last_modified 作為代理時間戳）。",
    "作答串解析：dp001_prac 的 binary_res 為以「@XX@」分隔之 0/1 字串，解析為題數與答對數。",
    "測驗時長之 ISO8601 格式（如 PT1M35S 表 1 分 35 秒）轉換為秒數。",
    "彙整合併：以使用者編號為鍵串接 user_data，分別彙總出事件層（edu_activity，101,680 筆）、"
    "使用者層（edu_users，313 筆）、影片瀏覽層（edu_video，4,567 筆）與題項難點層"
    "（edu_difficulty，1,264 筆）四個分析就緒表。",
]

PREP_CODE = '''import pandas as pd, re

def parse_binary_res(s):                      # '1@XX@0@XX@' -> (題數, 答對數)
    items = [x for x in str(s).split("@XX@") if x in ("0", "1")]
    return len(items), sum(map(int, items))

def parse_iso_duration(s):                    # 'PT1M35S' -> 95 秒
    m = re.fullmatch(r"PT(?:(\\d+)H)?(?:(\\d+)M)?(?:(\\d+)S)?", str(s))
    h, mi, sec = (int(g) if g else 0 for g in m.groups())
    return h * 3600 + mi * 60 + sec

users = pd.read_csv(SRC / "user_data.csv", sep="\\t").rename(columns=USER_RENAME)
# 成績缺漏：行為彙總保留全部 313 人、成績分析逐科排除
out[count_cols] = out[count_cols].fillna(0).astype(int)'''

THEMES = [
    ("報表1　學習平臺使用概況", "report1.png", "theme1"),
    ("報表2　學習行為與成績關聯", "report2.png", "theme2"),
    ("報表3　影片學習行為解析", "report3.png", "theme3"),
    ("報表4　科目與能力難點", "report4.png", "theme4"),
]


def style_run(run, size=None, bold=None):
    run.font.name = "Times New Roman"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_doc_font(doc):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")


def heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        style_run(run)
    return h


def para(doc, text, indent=True):
    return doc.add_paragraph(("　　" if indent else "") + text)


def add_toc(doc):
    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-2" \\h \\z \\u')
    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "（請在 Word 中按 F9 更新目錄）"
    run.append(t)
    fld.append(run)
    p._p.append(fld)


def simple_table(doc, headers, rows):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
    for r in rows:
        cells = tbl.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = str(v)


def code_block(doc, code):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def main():
    cover = json.loads((REPORT / "cover.json").read_text(encoding="utf-8"))
    findings = json.loads((REPORT / "findings.json").read_text(encoding="utf-8"))

    doc = Document()
    set_doc_font(doc)

    # ---- 封面 ----
    cover_lines = [
        ("「商業智慧」期末報告", 20, 60, True),
        (cover["title"], 24, 24, True),
        (f"學生：{cover['department']}　{cover['student_id']}　{cover['name']}", 14, 60, False),
    ]
    if cover.get("advisor"):
        cover_lines.append((f"指導教授：{cover['advisor']}", 14, 6, False))
    cover_lines.append((cover["date"], 14, 6, False))
    cover_lines.append((f"影片連結：{cover['video_url'] or '（待補）'}", 12, 12, False))
    for txt, size, before, bold in cover_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(before)
        run = p.add_run(txt)
        style_run(run, size=size, bold=bold)
    doc.add_page_break()

    heading(doc, "目錄", 1)
    add_toc(doc)
    doc.add_page_break()

    # ---- 一、分析主題 ----
    heading(doc, "一、分析主題", 1)
    para(doc, "本報告以「2025 教育大數據微學程教學用開放資料」為對象，"
              "分析 313 位國小學生於四個數位學習平臺一學期（2024 年 9 月至 2025 年 1 月）的"
              "行為紀錄，預計回答以下四個問題：")
    for i, q in enumerate(findings["research_questions"], 1):
        para(doc, f"{i}. {q}", indent=False)

    # ---- 二、資料來源與前處理 ----
    heading(doc, "二、資料來源與前處理", 1)
    heading(doc, "（一）資料來源", 2)
    para(doc, "資料取自教育大數據分析計畫辦公室之「2025 教育大數據微學程教學用開放資料」，"
              "共 11 個檔案（tab 分隔 CSV），涵蓋 313 位使用者在 dp001（影音學習）、dp002（測驗）、"
              "dp003（遊戲學習）、dp004（綜合學習）四個平臺的操作紀錄與國文、數學、英語測驗成績。"
              "本資料經匿名化與資料干擾處理，僅供教學用途，分析結果不得用於實務現象詮釋；"
              "本報告使用時已註明出處。")
    para(doc, "表 1　原始資料檔清單", indent=False)
    simple_table(doc, ["檔名", "筆數", "欄數", "內容"], FILES_TABLE)
    heading(doc, "（二）前處理步驟", 2)
    para(doc, "前處理以 Python（pandas）完成，步驟如下：", indent=False)
    for i, s in enumerate(PREP_STEPS, 1):
        para(doc, f"{i}. {s}", indent=False)
    para(doc, "表 2　欄位重新命名對照（節錄）", indent=False)
    simple_table(doc, ["原欄名", "更名後"], RENAME_TABLE)
    para(doc, "關鍵程式片段：", indent=False)
    code_block(doc, PREP_CODE)

    # ---- 三、資料分析與視覺化 ----
    heading(doc, "三、資料分析與視覺化", 1)
    para(doc, "前處理產出之四個分析就緒表，分別於 Power BI 服務建立語意模型與互動報表，"
              "以下依四大主題呈現儀表板與分析發現。")
    for idx, (title, png, key) in enumerate(THEMES, 1):
        heading(doc, f"（{'一二三四'[idx-1]}）{title}", 2)
        img = ASSETS / png
        if img.exists():
            doc.add_picture(str(img), width=Cm(16))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph(f"圖 {idx}　{title}")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for paragraph in findings[key]:
            para(doc, paragraph)

    # ---- 四、結論 ----
    heading(doc, "四、結論與建議", 1)
    for paragraph in findings["conclusion"]:
        para(doc, paragraph)

    # ---- 五、影片說明 ----
    heading(doc, "五、影片說明", 1)
    para(doc, "本報告輔以儀表板實際操作之解說影片，依序說明四大主題之分析發現與圖表互動方式。")
    para(doc, f"影片連結：{cover['video_url'] or '（待補）'}", indent=False)

    out = REPORT / "期末報告.docx"
    doc.save(out)
    print(f"saved: {out} ({out.stat().st_size//1024} KB)")


if __name__ == "__main__":
    main()
